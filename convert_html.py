from flask import Flask, request, send_file
from flask_cors import CORS
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import logging

# Configurar logging
logging.basicConfig(level=logging.INFO)
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')


app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:4200"}})

@app.route('/pdf', methods=['POST'])
def convert_html_to_pdf():
    content = request.json.get('html_content', '')

    if not content.strip():
        return {"error": "El contenido HTML está vacío."}, 400

    # Ajustar el contenido eliminando el uso de f-strings
    adjusted_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {
                size: 22cm 29.7cm; /* Tamaño personalizado */
                margin: 1cm;
            }
            body {
                font-family: Arial, sans-serif;
                font-size: 12px;
                line-height: 1.5;
                margin: 0;
                padding: 1cm;
                word-wrap: break-word;
            }
            img {
                max-width: 100%;
                height: auto;
            }
            div[style*="width: 21cm"] {
                width: auto !important;
                min-width: 0 !important;
            }
        </style>
    </head>
    <body>
    """ + content.replace("width: 21cm;", "") + """
    </body>
    </html>
    """

    pdf_file = BytesIO()
    try:
        html = HTML(string=adjusted_content, base_url="/")
        html.write_pdf(pdf_file)
        pdf_file.seek(0)
    except Exception as e:
        return {"error": f"Error al generar el PDF: {str(e)}"}, 500

    return send_file(
        pdf_file,
        mimetype='application/pdf',
        as_attachment=True,
        download_name='document.pdf'
    )
@app.route('/pdf-word', methods=['POST'])
def convert_html_to_word():
    try:
        logging.info("Inicio de la conversión de HTML a Word.")

        content = request.json.get('html_content', '')

        if not content.strip():
            logging.error("El contenido HTML está vacío.")
            return {"error": "El contenido HTML está vacío."}, 400

        logging.info("Procesando el contenido HTML con BeautifulSoup.")

        # Procesar el contenido HTML con BeautifulSoup
        soup = BeautifulSoup(content, "html.parser")
        document = Document()

        logging.info("Iterando sobre los elementos HTML y generando el documento Word.")

        for element in soup.find_all(['h1', 'h2', 'h3', 'p']):
            if element.name == 'h1':
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.bold = True
                run.font.size = Pt(24)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif element.name == 'h2':
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.bold = True
                run.font.size = Pt(18)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif element.name == 'h3':
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.italic = True
                run.bold = True
                run.font.size = Pt(14)
            elif element.name == 'p':
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.font.size = Pt(12)

        logging.info("Creación del documento Word completada.")

        # Guardar el documento en un archivo temporal
        word_file = BytesIO()
        document.save(word_file)
        word_file.seek(0)

        logging.info("El archivo Word fue creado exitosamente y está listo para ser enviado.")

        return send_file(
            word_file,
            as_attachment=True,
            download_name='document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logging.error(f"Error durante la conversión de HTML a Word: {str(e)}", exc_info=True)
        return {"error": f"Error al convertir HTML a Word: {str(e)}"}, 500


if __name__ == '__main__':
    app.run(debug=True)